"""
export_generator.py
Generate branded PDF and DOCX outputs that mirror the MedDoc AI web UI design.
"""
import io
import re
from datetime import datetime

# ── ReportLab ──────────────────────────────────────────────
from reportlab.lib.pagesizes import A4
from reportlab.lib.units import mm
from reportlab.lib.colors import HexColor, white
from reportlab.lib.styles import ParagraphStyle
from reportlab.lib.enums import TA_LEFT, TA_CENTER
from reportlab.platypus import (
    SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle,
    HRFlowable, KeepTogether,
)

# ── python-docx ────────────────────────────────────────────
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ── Brand palette ──────────────────────────────────────────
BLUE      = HexColor("#2563eb")
BLUE_L    = HexColor("#eff6ff")
BLUE_M    = HexColor("#dbeafe")
BLUE_D    = HexColor("#1d4ed8")
PURPLE    = HexColor("#9333ea")
PURPLE_L  = HexColor("#faf5ff")
GREEN     = HexColor("#16a34a")
GREEN_L   = HexColor("#f0fdf4")
SLATE_900 = HexColor("#0f172a")
SLATE_800 = HexColor("#1e293b")
SLATE_700 = HexColor("#334155")
SLATE_500 = HexColor("#64748b")
SLATE_400 = HexColor("#94a3b8")
SLATE_200 = HexColor("#e2e8f0")
SLATE_100 = HexColor("#f1f5f9")
SLATE_50  = HexColor("#f8fafc")

# DOCX equivalents (RGBColor)
def _rgb(hex_str: str):
    h = hex_str.lstrip("#")
    return RGBColor(int(h[0:2],16), int(h[2:4],16), int(h[4:6],16))

DOCX_BLUE     = _rgb("2563eb")
DOCX_BLUE_L   = _rgb("dbeafe")
DOCX_PURPLE   = _rgb("9333ea")
DOCX_GREEN    = _rgb("16a34a")
DOCX_SLATE_7  = _rgb("334155")
DOCX_SLATE_4  = _rgb("94a3b8")
DOCX_SLATE_9  = _rgb("0f172a")
DOCX_WHITE    = _rgb("ffffff")

MODE_META = {
    "summary": (BLUE,   BLUE_L,   "Clinical Summary"),
    "mask":    (PURPLE, PURPLE_L, "Privacy-Masked Document"),
    "report":  (GREEN,  GREEN_L,  "Structured Medical Report"),
}


def _flatten_lab(item) -> list[dict]:
    """Normalize a lab item (dict or string) into a list of flat row dicts."""
    if isinstance(item, str):
        return [{"name": item, "value": "", "unit": "", "range": ""}]
    if not isinstance(item, dict):
        return [{"name": str(item), "value": "", "unit": "", "range": ""}]

    name  = str(item.get("test_name") or item.get("name") or item.get("test") or "")
    value = str(item.get("value") or item.get("observed_value") or item.get("result") or "")
    unit  = str(item.get("unit") or item.get("units") or "")
    ref   = str(item.get("reference_range") or item.get("normal_range")
                or item.get("reference_interval") or item.get("range") or "")

    rows = [{"name": name, "value": value, "unit": unit, "range": ref}]

    # Expand nested sub_tests / results arrays
    for sub_key in ("sub_tests", "results"):
        subs = item.get(sub_key)
        if isinstance(subs, list):
            for sub in subs:
                rows.extend(_flatten_lab(sub))

    return rows

PAGE_W, PAGE_H = A4
MARGIN = 18 * mm

# ════════════════════════════════════════════════════════════
#  PDF GENERATION
# ════════════════════════════════════════════════════════════

def _pdf_styles(accent):
    def s(name, **kw):
        base = dict(fontName="Helvetica", fontSize=10, leading=14,
                    textColor=SLATE_700, spaceAfter=4)
        base.update(kw)
        return ParagraphStyle(name, **base)

    return {
        "h1":    s("h1",  fontName="Helvetica-Bold", fontSize=15, leading=20, textColor=white),
        "h2":    s("h2",  fontName="Helvetica-Bold", fontSize=11, leading=16, textColor=accent,
                   spaceBefore=10, spaceAfter=4),
        "h3":    s("h3",  fontName="Helvetica-Bold", fontSize=10, leading=14, textColor=SLATE_800,
                   spaceBefore=8,  spaceAfter=3),
        "body":  s("body",fontSize=9.5, leading=14, textColor=SLATE_700, spaceAfter=3),
        "mono":  s("mono",fontName="Courier", fontSize=8.5, leading=13, textColor=SLATE_700, spaceAfter=2),
        "meta":  s("meta",fontSize=8, leading=12, textColor=SLATE_400),
        "label": s("label",fontName="Helvetica-Bold", fontSize=7.5, leading=11, textColor=SLATE_400,
                   spaceAfter=1),
        "value": s("value",fontSize=9, leading=13, textColor=SLATE_800),
        "th":    s("th",  fontName="Helvetica-Bold", fontSize=8.5, leading=12, textColor=white),
        "td":    s("td",  fontSize=8.5, leading=13, textColor=SLATE_700),
    }


def _pdf_header(job_data: dict, accent, mode_label: str, st: dict) -> Table:
    w = PAGE_W - 2 * MARGIN
    ts = datetime.fromisoformat(job_data["timestamp"].replace("Z", ""))
    date_str = ts.strftime("%d %b %Y, %H:%M UTC")

    title = Paragraph(
        f'<font name="Helvetica-Bold" size="15" color="white">MedDoc AI</font>'
        f'<font name="Helvetica" size="10" color="#bfdbfe">  —  {mode_label}</font>',
        st["h1"],
    )
    meta = Paragraph(
        f'<font color="#93c5fd">Date:</font> <font color="white">{date_str}</font>',
        ParagraphStyle("hm", fontName="Helvetica", fontSize=8.5, leading=12, textColor=white),
    )

    tbl = Table([[title], [meta]], colWidths=[w], rowHeights=[20*mm, 11*mm])
    tbl.setStyle(TableStyle([
        ("BACKGROUND",   (0,0), (-1,-1), accent),
        ("LEFTPADDING",  (0,0), (-1,-1), 14),
        ("RIGHTPADDING", (0,0), (-1,-1), 14),
        ("TOPPADDING",   (0,0), (0, 0),  9),
        ("BOTTOMPADDING",(0,0), (0, 0),  2),
        ("TOPPADDING",   (0,1), (0, 1),  2),
        ("BOTTOMPADDING",(0,1), (0, 1),  9),
    ]))
    return tbl


def _md_to_pdf_elements(text: str, st: dict) -> list:
    els = []
    for line in text.split("\n"):
        l = line.strip()
        if not l:
            els.append(Spacer(1, 3))
            continue
        safe = l.replace("&","&amp;").replace("<","&lt;").replace(">","&gt;")
        safe = re.sub(r"\*\*(.+?)\*\*", r"<b>\1</b>", safe)
        if l.startswith("## ") or l.startswith("# "):
            els.append(Paragraph(re.sub(r"^#+\s*","",l), st["h2"]))
        elif l.startswith("### "):
            els.append(Paragraph(l[4:], st["h3"]))
        elif l.startswith(("- ","* ")):
            els.append(Paragraph("•  " + safe[2:], st["body"]))
        else:
            els.append(Paragraph(safe, st["body"]))
    return els


def _pdf_report_section(story, data: dict, st: dict, accent):
    w = PAGE_W - 2 * MARGIN
    SCALARS = [
        ("patient_name","Patient Name"), ("patient_age","Age"),
        ("patient_gender","Gender"),     ("visit_date","Visit Date"),
        ("doctor_name","Doctor"),        ("facility_name","Facility"),
        ("chief_complaint","Chief Complaint"), ("document_type","Document Type"),
    ]
    # 2-col scalar grid
    grid_rows = []
    pairs = [SCALARS[i:i+2] for i in range(0, len(SCALARS), 2)]
    for pair in pairs:
        row = []
        for key, label in pair:
            val = str(data.get(key) or "—")
            row.append(Paragraph(label.upper(), st["label"]))
            row.append(Paragraph(val, st["value"]))
        if len(row) == 4:
            pass
        else:
            row += [Paragraph("","label"), Paragraph("","value")]
        grid_rows.append(row)

    if grid_rows:
        half = (w - 6*mm) / 2
        tbl = Table(grid_rows, colWidths=[half*0.38, half*0.62, half*0.38, half*0.62])
        tbl.setStyle(TableStyle([
            ("BACKGROUND",(0,0),(0,-1), SLATE_100),
            ("BACKGROUND",(2,0),(2,-1), SLATE_100),
            ("GRID",(0,0),(-1,-1), 0.5, SLATE_200),
            ("LEFTPADDING",(0,0),(-1,-1), 7),
            ("RIGHTPADDING",(0,0),(-1,-1), 7),
            ("TOPPADDING",(0,0),(-1,-1), 5),
            ("BOTTOMPADDING",(0,0),(-1,-1), 5),
            ("VALIGN",(0,0),(-1,-1),"TOP"),
        ]))
        story.append(tbl)
        story.append(Spacer(1, 4*mm))

    # Diagnosis
    diag = data.get("diagnosis") or []
    if diag:
        story.append(Paragraph("Diagnosis", st["h3"]))
        for d in diag:
            story.append(Paragraph(f"•  {d}", st["body"]))
        story.append(Spacer(1, 3*mm))

    # Medications table
    meds = data.get("medications") or []
    if meds:
        story.append(Paragraph("Medications", st["h3"]))
        rows = [[Paragraph(h, st["th"]) for h in ["Medication","Dosage","Frequency","Duration"]]]
        for m in meds:
            rows.append([Paragraph(str(m.get(k) or "—"), st["td"])
                         for k in ("name","dosage","frequency","duration")])
        cw = w / 4
        tbl = Table(rows, colWidths=[cw]*4)
        tbl.setStyle(TableStyle([
            ("BACKGROUND",(0,0),(-1,0), accent),
            ("GRID",(0,0),(-1,-1), 0.5, SLATE_200),
            ("ROWBACKGROUNDS",(0,1),(-1,-1), [white, SLATE_50]),
            ("LEFTPADDING",(0,0),(-1,-1), 7),
            ("RIGHTPADDING",(0,0),(-1,-1), 7),
            ("TOPPADDING",(0,0),(-1,-1), 5),
            ("BOTTOMPADDING",(0,0),(-1,-1), 5),
            ("VALIGN",(0,0),(-1,-1),"MIDDLE"),
        ]))
        story.append(tbl)
        story.append(Spacer(1, 3*mm))

    # Lab tests
    labs = data.get("lab_tests") or []
    if labs:
        story.append(Paragraph("Lab Tests", st["h3"]))
        flat_rows = []
        for lab in labs:
            flat_rows.extend(_flatten_lab(lab))
        if flat_rows:
            headers = [Paragraph(h, st["th"]) for h in ["Test Name", "Value", "Unit", "Reference Range"]]
            tbl_rows = [headers]
            for r in flat_rows:
                tbl_rows.append([
                    Paragraph(r["name"],  st["td"]),
                    Paragraph(r["value"], st["td"]),
                    Paragraph(r["unit"],  st["td"]),
                    Paragraph(r["range"], st["td"]),
                ])
            col_w = [w * 0.40, w * 0.20, w * 0.15, w * 0.25]
            lab_tbl = Table(tbl_rows, colWidths=col_w)
            lab_tbl.setStyle(TableStyle([
                ("BACKGROUND",   (0, 0), (-1, 0), GREEN),
                ("GRID",         (0, 0), (-1,-1), 0.5, SLATE_200),
                ("ROWBACKGROUNDS",(0, 1),(-1,-1), [white, SLATE_50]),
                ("LEFTPADDING",  (0, 0), (-1,-1), 7),
                ("RIGHTPADDING", (0, 0), (-1,-1), 7),
                ("TOPPADDING",   (0, 0), (-1,-1), 5),
                ("BOTTOMPADDING",(0, 0), (-1,-1), 5),
                ("VALIGN",       (0, 0), (-1,-1), "MIDDLE"),
            ]))
            story.append(lab_tbl)
        story.append(Spacer(1, 3*mm))

    # Vitals
    vitals = data.get("vitals") or {}
    if isinstance(vitals, dict) and vitals:
        story.append(Paragraph("Vitals", st["h3"]))
        vrows = [[Paragraph(k, st["label"]), Paragraph(str(v), st["value"])]
                 for k,v in vitals.items()]
        vw = w * 0.6
        tbl = Table(vrows, colWidths=[vw*0.4, vw*0.6])
        tbl.setStyle(TableStyle([
            ("BACKGROUND",(0,0),(0,-1), SLATE_100),
            ("GRID",(0,0),(-1,-1), 0.5, SLATE_200),
            ("LEFTPADDING",(0,0),(-1,-1),7),
            ("TOPPADDING",(0,0),(-1,-1),4),
            ("BOTTOMPADDING",(0,0),(-1,-1),4),
        ]))
        story.append(tbl)
        story.append(Spacer(1, 3*mm))

    for key, label in [("doctor_notes","Doctor Notes"), ("follow_up","Follow-up Recommendations")]:
        val = data.get(key,"")
        if val:
            story.append(Paragraph(label, st["h3"]))
            story.append(Paragraph(str(val), st["body"]))
            story.append(Spacer(1, 2*mm))


def _pdf_entities(story, entities: dict, st: dict):
    if not entities:
        return
    story.append(Spacer(1, 5*mm))
    story.append(HRFlowable(width="100%", thickness=0.5, color=SLATE_200))
    story.append(Spacer(1, 4*mm))
    story.append(Paragraph("Detected Entities (Rule-Based Scanner)", st["h2"]))
    w = PAGE_W - 2 * MARGIN
    rows = []
    for key, vals in entities.items():
        rows.append([
            Paragraph(key.upper(), st["label"]),
            Paragraph(", ".join(str(v) for v in vals), st["value"]),
        ])
    tbl = Table(rows, colWidths=[w*0.25, w*0.75])
    tbl.setStyle(TableStyle([
        ("BACKGROUND",(0,0),(0,-1), SLATE_100),
        ("GRID",(0,0),(-1,-1), 0.5, SLATE_200),
        ("LEFTPADDING",(0,0),(-1,-1),8),
        ("TOPPADDING",(0,0),(-1,-1),5),
        ("BOTTOMPADDING",(0,0),(-1,-1),5),
        ("VALIGN",(0,0),(-1,-1),"TOP"),
    ]))
    story.append(tbl)


def generate_pdf(job_data: dict) -> bytes:
    mode = job_data.get("mode", "summary")
    accent, _, mode_label = MODE_META.get(mode, MODE_META["summary"])
    st = _pdf_styles(accent)

    buf = io.BytesIO()
    doc = SimpleDocTemplate(
        buf, pagesize=A4,
        leftMargin=MARGIN, rightMargin=MARGIN,
        topMargin=MARGIN, bottomMargin=MARGIN,
        title=f"MedDoc AI — {mode_label}",
        author="MedDoc AI",
    )

    story = [_pdf_header(job_data, accent, mode_label, st), Spacer(1, 6*mm)]

    output  = job_data.get("output", {})
    content = output.get("content", "")

    if mode == "summary":
        story.extend(_md_to_pdf_elements(str(content), st))

    elif mode == "mask":
        story.append(Paragraph("Masked Document", st["h2"]))
        story.append(HRFlowable(width="100%", thickness=0.5, color=SLATE_200))
        story.append(Spacer(1, 3*mm))
        safe = str(content).replace("&","&amp;").replace("<","&lt;").replace(">","&gt;")
        safe = re.sub(
            r"\[(PATIENT_NAME|DOCTOR_NAME|FACILITY_NAME|PHONE|EMAIL|ADDRESS|DOB|ID_NUMBER|REDACTED|AADHAR|PAN)\]",
            r'<font color="#9333ea" name="Helvetica-Bold">[\1]</font>', safe,
        )
        for line in safe.split("\n"):
            if line.strip():
                story.append(Paragraph(line, st["mono"]))
            else:
                story.append(Spacer(1, 2))

    elif mode == "report":
        if isinstance(content, dict) and not content.get("parse_error"):
            _pdf_report_section(story, content, st, accent)
        else:
            raw = content.get("raw", str(content)) if isinstance(content, dict) else str(content)
            story.append(Paragraph(raw, st["mono"]))

    _pdf_entities(story, job_data.get("entities", {}), st)

    # Footer via canvas
    def _footer(canvas, doc):
        canvas.saveState()
        canvas.setFont("Helvetica", 7.5)
        canvas.setFillColor(SLATE_400)
        canvas.drawString(MARGIN, 12*mm, "MedDoc AI — Privacy-First Medical Document Processor")
        canvas.drawRightString(PAGE_W - MARGIN, 12*mm, f"Page {doc.page}")
        canvas.restoreState()

    doc.build(story, onFirstPage=_footer, onLaterPages=_footer)
    return buf.getvalue()


# ════════════════════════════════════════════════════════════
#  DOCX GENERATION
# ════════════════════════════════════════════════════════════

def _docx_set_cell_bg(cell, hex_color: str):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color.lstrip("#"))
    tcPr.append(shd)


def _docx_header(doc: Document, job_data: dict, mode_label: str, mode: str):
    accent_hex = {"summary":"2563eb","mask":"9333ea","report":"16a34a"}.get(mode,"2563eb")

    # Title row
    tbl = doc.add_table(rows=2, cols=1)
    tbl.style = "Table Grid"

    r0 = tbl.rows[0].cells[0]
    _docx_set_cell_bg(r0, accent_hex)
    p = r0.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run("MedDoc AI")
    run.bold = True; run.font.size = Pt(16); run.font.color.rgb = DOCX_WHITE
    sub = p.add_run(f"  —  {mode_label}")
    sub.font.size = Pt(10); sub.font.color.rgb = _rgb("bfdbfe")

    r1 = tbl.rows[1].cells[0]
    _docx_set_cell_bg(r1, accent_hex)
    ts = datetime.fromisoformat(job_data["timestamp"].replace("Z",""))
    p2 = r1.paragraphs[0]
    p2.alignment = WD_ALIGN_PARAGRAPH.LEFT
    lbl = p2.add_run("Date: "); lbl.font.size = Pt(8.5); lbl.font.color.rgb = _rgb("93c5fd")
    v   = p2.add_run(ts.strftime("%d %b %Y, %H:%M UTC")); v.font.size = Pt(8.5); v.font.color.rgb = DOCX_WHITE

    # Remove table borders
    for row in tbl.rows:
        for cell in row.cells:
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            tcBorders = OxmlElement("w:tcBorders")
            for side in ("top","left","bottom","right","insideH","insideV"):
                border = OxmlElement(f"w:{side}")
                border.set(qn("w:val"), "none")
                tcBorders.append(border)
            tcPr.append(tcBorders)

    doc.add_paragraph()


def _docx_heading(doc, text: str, level: int, color: RGBColor):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(13 - level)
    run.font.color.rgb = color
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after  = Pt(4)
    return p


def _docx_hr(doc):
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"),  "single")
    bottom.set(qn("w:sz"),   "4")
    bottom.set(qn("w:space"),"1")
    bottom.set(qn("w:color"),"e2e8f0")
    pBdr.append(bottom)
    pPr.append(pBdr)


def _docx_md_content(doc: Document, text: str, accent: RGBColor):
    for line in text.split("\n"):
        l = line.strip()
        if not l:
            doc.add_paragraph()
            continue
        if l.startswith("## ") or l.startswith("# "):
            _docx_heading(doc, re.sub(r"^#+\s*","",l), 2, accent)
        elif l.startswith("### "):
            _docx_heading(doc, l[4:], 3, DOCX_SLATE_9)
        elif l.startswith(("- ","* ")):
            p = doc.add_paragraph(style="List Bullet")
            run = p.add_run(re.sub(r"\*\*(.+?)\*\*", r"\1", l[2:]))
            run.font.size = Pt(10); run.font.color.rgb = DOCX_SLATE_7
        else:
            p = doc.add_paragraph()
            cleaned = re.sub(r"\*\*(.+?)\*\*", r"\1", l)
            run = p.add_run(cleaned)
            run.font.size = Pt(10); run.font.color.rgb = DOCX_SLATE_7


def _docx_report_section(doc: Document, data: dict, accent: RGBColor):
    SCALARS = [
        ("patient_name","Patient Name"), ("patient_age","Age"),
        ("patient_gender","Gender"),     ("visit_date","Visit Date"),
        ("doctor_name","Doctor"),        ("facility_name","Facility"),
        ("chief_complaint","Chief Complaint"), ("document_type","Document Type"),
    ]
    # Scalar table (2 per row)
    tbl = doc.add_table(rows=0, cols=4)
    tbl.style = "Table Grid"
    pairs = [SCALARS[i:i+2] for i in range(0, len(SCALARS), 2)]
    for pair in pairs:
        row = tbl.add_row().cells
        for idx, (key, label) in enumerate(pair):
            _docx_set_cell_bg(row[idx*2], "f1f5f9")
            lp = row[idx*2].paragraphs[0]
            lr = lp.add_run(label.upper())
            lr.bold = True; lr.font.size = Pt(7.5); lr.font.color.rgb = DOCX_SLATE_4

            val = str(data.get(key) or "—")
            vp = row[idx*2+1].paragraphs[0]
            vr = vp.add_run(val)
            vr.font.size = Pt(9); vr.font.color.rgb = DOCX_SLATE_9

    doc.add_paragraph()

    # Diagnosis
    diag = data.get("diagnosis") or []
    if diag:
        _docx_heading(doc, "Diagnosis", 3, DOCX_SLATE_9)
        for d in diag:
            p = doc.add_paragraph(style="List Bullet")
            p.add_run(str(d)).font.size = Pt(10)

    # Medications
    meds = data.get("medications") or []
    if meds:
        _docx_heading(doc, "Medications", 3, DOCX_SLATE_9)
        med_tbl = doc.add_table(rows=1, cols=4)
        med_tbl.style = "Table Grid"
        hdr = med_tbl.rows[0].cells
        for i, h in enumerate(["Medication","Dosage","Frequency","Duration"]):
            _docx_set_cell_bg(hdr[i], "2563eb")
            p = hdr[i].paragraphs[0]
            r = p.add_run(h); r.bold = True; r.font.size = Pt(9); r.font.color.rgb = DOCX_WHITE
        for idx, m in enumerate(meds):
            row = med_tbl.add_row().cells
            bg = "f8fafc" if idx % 2 == 0 else "f1f5f9"
            for i, k in enumerate(("name","dosage","frequency","duration")):
                _docx_set_cell_bg(row[i], bg)
                row[i].paragraphs[0].add_run(str(m.get(k) or "—")).font.size = Pt(9)
        doc.add_paragraph()

    # Lab tests
    labs = data.get("lab_tests") or []
    if labs:
        _docx_heading(doc, "Lab Tests", 3, DOCX_SLATE_9)
        flat_rows = []
        for lab in labs:
            flat_rows.extend(_flatten_lab(lab))
        if flat_rows:
            lab_tbl = doc.add_table(rows=1, cols=4)
            lab_tbl.style = "Table Grid"
            hdr = lab_tbl.rows[0].cells
            for i, h in enumerate(["Test Name", "Value", "Unit", "Reference Range"]):
                _docx_set_cell_bg(hdr[i], "16a34a")
                p = hdr[i].paragraphs[0]
                r = p.add_run(h); r.bold = True; r.font.size = Pt(9); r.font.color.rgb = DOCX_WHITE
            for idx, r in enumerate(flat_rows):
                row = lab_tbl.add_row().cells
                bg = "f8fafc" if idx % 2 == 0 else "f1f5f9"
                for i, key in enumerate(("name", "value", "unit", "range")):
                    _docx_set_cell_bg(row[i], bg)
                    row[i].paragraphs[0].add_run(str(r[key])).font.size = Pt(9)
        doc.add_paragraph()

    # Vitals
    vitals = data.get("vitals") or {}
    if isinstance(vitals, dict) and vitals:
        _docx_heading(doc, "Vitals", 3, DOCX_SLATE_9)
        vtbl = doc.add_table(rows=0, cols=2)
        vtbl.style = "Table Grid"
        for k, v in vitals.items():
            row = vtbl.add_row().cells
            _docx_set_cell_bg(row[0], "f1f5f9")
            row[0].paragraphs[0].add_run(str(k)).font.size = Pt(9)
            row[1].paragraphs[0].add_run(str(v)).font.size = Pt(9)
        doc.add_paragraph()

    for key, label in [("doctor_notes","Doctor Notes"),("follow_up","Follow-up")]:
        val = data.get(key,"")
        if val:
            _docx_heading(doc, label, 3, DOCX_SLATE_9)
            p = doc.add_paragraph(); p.add_run(str(val)).font.size = Pt(10)


def _docx_entities(doc: Document, entities: dict):
    if not entities:
        return
    _docx_hr(doc)
    _docx_heading(doc, "Detected Entities (Rule-Based Scanner)", 2, DOCX_SLATE_9)
    tbl = doc.add_table(rows=1, cols=2)
    tbl.style = "Table Grid"
    hdr = tbl.rows[0].cells
    _docx_set_cell_bg(hdr[0], "f1f5f9"); hdr[0].paragraphs[0].add_run("TYPE").font.bold = True
    _docx_set_cell_bg(hdr[1], "f1f5f9"); hdr[1].paragraphs[0].add_run("VALUES").font.bold = True
    for key, vals in entities.items():
        row = tbl.add_row().cells
        row[0].paragraphs[0].add_run(key.upper()).font.size = Pt(9)
        row[1].paragraphs[0].add_run(", ".join(str(v) for v in vals)).font.size = Pt(9)


def generate_docx(job_data: dict) -> bytes:
    mode = job_data.get("mode","summary")
    accent_hex = {"summary":"2563eb","mask":"9333ea","report":"16a34a"}.get(mode,"2563eb")
    accent_rgb = _rgb(accent_hex)
    _, _, mode_label = MODE_META.get(mode, MODE_META["summary"])

    doc = Document()
    for section in doc.sections:
        section.top_margin    = Cm(1.8)
        section.bottom_margin = Cm(1.8)
        section.left_margin   = Cm(2.2)
        section.right_margin  = Cm(2.2)

    # Set default font
    doc.styles["Normal"].font.name = "Calibri"
    doc.styles["Normal"].font.size = Pt(10)

    _docx_header(doc, job_data, mode_label, mode)

    output  = job_data.get("output", {})
    content = output.get("content", "")

    if mode == "summary":
        _docx_md_content(doc, str(content), accent_rgb)

    elif mode == "mask":
        _docx_heading(doc, "Masked Document", 2, accent_rgb)
        _docx_hr(doc)
        for line in str(content).split("\n"):
            if not line.strip():
                doc.add_paragraph()
                continue
            p = doc.add_paragraph()
            # Split on mask tokens and style them
            parts = re.split(r"(\[[A-Z_]+\])", line)
            for part in parts:
                run = p.add_run(part)
                run.font.name = "Courier New"
                run.font.size = Pt(9)
                if re.match(r"^\[[A-Z_]+\]$", part):
                    run.bold = True
                    run.font.color.rgb = DOCX_PURPLE
                else:
                    run.font.color.rgb = DOCX_SLATE_7

    elif mode == "report":
        if isinstance(content, dict) and not content.get("parse_error"):
            _docx_report_section(doc, content, accent_rgb)
        else:
            raw = content.get("raw", str(content)) if isinstance(content, dict) else str(content)
            p = doc.add_paragraph()
            r = p.add_run(raw); r.font.name = "Courier New"; r.font.size = Pt(9)

    _docx_entities(doc, job_data.get("entities", {}))

    # Footer
    section = doc.sections[0]
    footer = section.footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = fp.add_run("MedDoc AI — Privacy-First Medical Document Processor")
    r.font.size = Pt(8); r.font.color.rgb = DOCX_SLATE_4

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
